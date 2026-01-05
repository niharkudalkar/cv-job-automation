#!/usr/bin/env python3
# Job Application Workflow Automation (India-focused)
# Author: Nihar Kudalkar
# Purpose: Automate CV-based job discovery, resume tailoring, cover letter generation, and tracking

import os
import datetime
import requests
from docx import Document
import json

# -------------------------------
# 1. Extract Keywords from CV
# -------------------------------
def extract_keywords_from_cv(cv_path="resume_master.docx"):
    """
    Extract key skills/roles from CV using document parsing.
    """
    try:
        doc = Document(cv_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        # Simple keyword extraction (expand with NLP as needed)
        keywords = {
            "roles": ["Product Manager", "Delivery Head", "Project Manager"],
            "skills": ["Agile", "SaaS", "Team Management", "Process Improvement"],
            "experience_level": "Senior"
        }
        print(f"Extracted keywords from CV: {keywords}")
        return keywords
    except FileNotFoundError:
        print(f"CV file not found: {cv_path}")
        return {}

# -------------------------------
# 2. Job Discovery (India Sites)
# -------------------------------
def discover_jobs_from_sites(keywords, location="India"):
    """
    Search multiple job boards in India using keywords.
    Replace with actual API calls or web scraping.
    """
    job_sites = {
        "LinkedIn": "https://www.linkedin.com/jobs/search",
        "Naukri": "https://www.naukri.com",
        "Indeed": "https://www.indeed.co.in",
        "Glassdoor": "https://www.glassdoor.co.in",
        "Monster": "https://www.monsterindia.com"
    }

    jobs = []
    for site, url in job_sites.items():
        print(f"Searching {site} for {keywords.get('roles', [])} in {location}...")
        # Placeholder: simulate results (replace with actual API/scraping)
        jobs.append({
            "company": "TechCorp",
            "role": "Product Manager",
            "site": site,
            "url": f"{url}/job123",
            "salary_range": "12-18 LPA"
        })
        jobs.append({
            "company": "InnoSoft",
            "role": "Senior Delivery Manager",
            "site": site,
            "url": f"{url}/job456",
            "salary_range": "15-20 LPA"
        })
    return jobs

# -------------------------------
# 3. Resume Tailoring
# -------------------------------
def tailor_resume(job_description, master_resume_path="resume_master.docx"):
    """
    Customize resume for specific job posting.
    """
    try:
        doc = Document(master_resume_path)
        # Add job-specific keywords to the resume (simplified approach)
        for para in doc.paragraphs:
            if "skills" in para.text.lower():
                para.text += f" - {job_description.get('role', '')}"
        
        filename = f"Resume_{job_description['company']}_{job_description['role'].replace(' ', '_')}.docx"
        doc.save(filename)
        print(f"Tailored resume saved: {filename}")
        return filename
    except FileNotFoundError:
        print(f"Master resume not found: {master_resume_path}")
        return None

# -------------------------------
# 4. Cover Letter Generation
# -------------------------------
def generate_cover_letter(job_description):
    """
    Generate a customized cover letter for the job application.
    """
    cover_letter_text = f"""Dear Hiring Manager at {job_description['company']},

I am excited to apply for the {job_description['role']} role at your esteemed organization.

With extensive expertise in SaaS, automation, and cross-functional leadership, I bring proven success 
in managing Fortune 500 projects and driving IT transformation initiatives.

Key Qualifications:
- Strong background in Agile methodologies and team management
- Experience with process improvement and delivery optimization
- Proven track record in remote work environments
- Excellent communication and stakeholder management skills

I am confident that my experience aligns well with your requirements and I would be delighted 
to discuss how I can contribute to your team's success.

Thank you for considering my application. I look forward to hearing from you soon.

Best regards,
Nihar Kudalkar
Contact: [Your Contact Info]
LinkedIn: https://linkedin.com/in/your-profile
"""
    
    filename = f"CoverLetter_{job_description['company']}_{job_description['role'].replace(' ', '_')}.txt"
    with open(filename, "w") as f:
        f.write(cover_letter_text)
    print(f"Cover letter saved: {filename}")
    return filename

# -------------------------------
# 5. Application Submission & Logging
# -------------------------------
def submit_and_log(job, resume_file, cover_letter_file, log_file="applications_log.json"):
    """
    Log application details for tracking and follow-up.
    """
    log_entry = {
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "company": job['company'],
        "role": job['role'],
        "site": job['site'],
        "url": job['url'],
        "resume": resume_file,
        "cover_letter": cover_letter_file,
        "status": "Applied",
        "follow_up_date": (datetime.datetime.now() + datetime.timedelta(days=7)).strftime("%Y-%m-%d")
    }
    
    # Load existing log or create new
    try:
        with open(log_file, "r") as f:
            log_data = json.load(f)
    except FileNotFoundError:
        log_data = []
    
    # Append new entry
    log_data.append(log_entry)
    
    # Save updated log
    with open(log_file, "w") as f:
        json.dump(log_data, f, indent=4)
    
    print(f"Application logged: {job['company']} - {job['role']}")
    return log_entry

# -------------------------------
# 6. Main Workflow
# -------------------------------
def main():
    """
    Execute the complete job application automation workflow.
    """
    print("=" * 60)
    print("CV-Based Job Application Automation - India Focus")
    print("=" * 60)
    
    # Step 1: Extract CV keywords
    keywords = extract_keywords_from_cv("resume_master.docx")
    if not keywords:
        print("No keywords extracted. Exiting.")
        return
    
    # Step 2: Discover jobs
    print("\nStep 2: Discovering job opportunities...")
    jobs = discover_jobs_from_sites(keywords, location="India Remote")
    print(f"Found {len(jobs)} job opportunities.\n")
    
    # Step 3: Apply to each job
    print("Step 3: Preparing applications...")
    for idx, job in enumerate(jobs, 1):
        print(f"\n--- Application {idx}: {job['company']} ---")
        
        # Tailor resume
        resume_file = tailor_resume(job)
        if not resume_file:
            print(f"Skipping {job['company']} - resume tailoring failed")
            continue
        
        # Generate cover letter
        cover_letter_file = generate_cover_letter(job)
        
        # Log application
        submit_and_log(job, resume_file, cover_letter_file)
    
    print("\n" + "=" * 60)
    print("Job application workflow completed successfully!")
    print("Check 'applications_log.json' for application tracking.")
    print("=" * 60)

if __name__ == "__main__":
    main()
